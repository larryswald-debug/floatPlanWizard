#!/usr/bin/env python3
"""Build the PDF, EPUB, DOCX, and standalone cover from one HTML manuscript."""

from __future__ import annotations

from copy import deepcopy
from html import escape
import io
import re
import shutil
from pathlib import Path
from zlib import crc32
from zipfile import ZIP_DEFLATED, ZIP_STORED, ZipFile

from lxml import etree, html
from PIL import Image as PILImage, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.fonts import addMapping
from reportlab.lib.pagesizes import inch
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Flowable,
    Frame,
    Image,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
)
from reportlab.platypus.tableofcontents import TableOfContents

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript" / "solo-boating-safety.html"
GRAPHICS = ROOT / "assets" / "graphics" / "export"
COVER_ART = ROOT / "assets" / "cover" / "cover-art-original.png"
EPUB_CSS = ROOT / "styles" / "epub.css"
DIST = ROOT / "dist"
WEB_DOWNLOADS = ROOT.parents[1] / "downloads"

PDF_PATH = DIST / "floatplanwizard-solo-boating-safety-guide.pdf"
EPUB_PATH = DIST / "solo-boating-safety-a-practical-guide.epub"
DOCX_PATH = DIST / "solo-boating-safety-kindle-source.docx"
COVER_PATH = DIST / "solo-boating-safety-cover.jpg"

TITLE = "Solo Boating Safety"
SUBTITLE = "A Practical Guide from Kayaks to Cruisers"
FULL_TITLE = f"{TITLE}: {SUBTITLE}"
AUTHOR = "Larry W."
PUBLISHER = "FloatPlanWizard"
DESCRIPTION = (
    "Practical preparation, communications, self-recovery, weather awareness, "
    "float planning, and pre-departure checklists for recreational boaters who go out alone."
)
KEYWORDS = "solo boating, boating safety, float plan, shore contact, PFD, marine VHF, self-recovery"

ARIAL = Path("/System/Library/Fonts/Supplemental/Arial.ttf")
ARIAL_BOLD = Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")
ARIAL_ITALIC = Path("/System/Library/Fonts/Supplemental/Arial Italic.ttf")


def parse_manuscript():
    return html.parse(str(MANUSCRIPT)).getroot()


def inner_markup(element) -> str:
    """Return ReportLab-compatible inline HTML for one source element."""
    pieces = [escape(element.text or "")]
    for child in element:
        tag = child.tag.lower() if isinstance(child.tag, str) else ""
        text = escape("".join(child.itertext()))
        if tag in {"strong", "b"}:
            pieces.append(f"<b>{text}</b>")
        elif tag in {"em", "i"}:
            pieces.append(f"<i>{text}</i>")
        elif tag == "a":
            pieces.append(f'<link href="{escape(child.get("href", ""), quote=True)}">{text}</link>')
        elif tag == "br":
            pieces.append("<br/>")
        else:
            pieces.append(text)
        pieces.append(escape(child.tail or ""))
    return "".join(pieces).strip()


def plain_text(element) -> str:
    """Return normalized visible text for labels, tooltips, and validation-friendly names."""
    return " ".join(" ".join(element.itertext()).split())


def checklist_field_name(checklist, item_index: int) -> str:
    """Return the stable cross-format field name for a Chapter 13 checklist item."""
    group = checklist.getparent()
    group_id = group.get("id", "checklist") if group is not None else "checklist"
    group_slug = re.sub(r"[^a-z0-9]+", "_", group_id.removeprefix("checklist-").lower()).strip("_")
    return f"ch13_{group_slug}_{item_index:02d}"


def crop_to_ratio(image: PILImage.Image, target_width: int, target_height: int) -> PILImage.Image:
    src_ratio = image.width / image.height
    target_ratio = target_width / target_height
    if src_ratio > target_ratio:
        width = round(image.height * target_ratio)
        left = (image.width - width) // 2
        image = image.crop((left, 0, left + width, image.height))
    elif src_ratio < target_ratio:
        height = round(image.width / target_ratio)
        top = (image.height - height) // 2
        image = image.crop((0, top, image.width, top + height))
    return image.resize((target_width, target_height), PILImage.Resampling.LANCZOS)


def text_with_shadow(draw: ImageDraw.ImageDraw, xy, text: str, font, *, fill="white", spacing=12,
                     anchor="la", align="left", shadow=3) -> None:
    x, y = xy
    draw.multiline_text(
        (x + shadow, y + shadow), text, font=font, fill=(0, 18, 32, 180),
        spacing=spacing, anchor=anchor, align=align,
    )
    draw.multiline_text(
        (x, y), text, font=font, fill=fill, spacing=spacing, anchor=anchor, align=align,
    )


def build_cover() -> Path:
    DIST.mkdir(parents=True, exist_ok=True)
    art = PILImage.open(COVER_ART).convert("RGB")
    cover = crop_to_ratio(art, 1600, 2560).convert("RGBA")

    shade = PILImage.new("RGBA", cover.size, (0, 0, 0, 0))
    shade_draw = ImageDraw.Draw(shade)
    for y in range(1900):
        alpha = max(0, round(196 * (1 - y / 1900)))
        shade_draw.line((0, y, 1600, y), fill=(3, 26, 45, alpha))
    cover = PILImage.alpha_composite(cover, shade)
    draw = ImageDraw.Draw(cover)

    title_font = ImageFont.truetype(str(ARIAL_BOLD), 132)
    subtitle_font = ImageFont.truetype(str(ARIAL), 54)
    topic_font = ImageFont.truetype(str(ARIAL_BOLD), 31)
    author_font = ImageFont.truetype(str(ARIAL_BOLD), 38)
    small_font = ImageFont.truetype(str(ARIAL), 27)

    draw.rounded_rectangle((102, 116, 536, 170), radius=17, fill=(8, 127, 140, 232))
    draw.text((319, 144), "FLOATPLANWIZARD", font=topic_font, fill="white", anchor="mm")
    text_with_shadow(draw, (104, 292), "SOLO BOATING\nSAFETY", title_font, spacing=-5)
    draw.rectangle((106, 594, 615, 607), fill=(230, 169, 61, 255))
    text_with_shadow(draw, (106, 675), "A Practical Guide\nfrom Kayaks to Cruisers", subtitle_font, spacing=10)
    text_with_shadow(
        draw,
        (108, 875),
        "Preparation  |  Communications  |  Self-Recovery\nWeather  |  Float Plans  |  Pre-Departure Checklists",
        topic_font,
        spacing=14,
    )

    draw.rounded_rectangle((90, 2286, 1510, 2496), radius=28, fill=(3, 26, 45, 214))
    draw.text((126, 2340), "LARRY W.", font=author_font, fill="white")
    draw.text((126, 2394), "Founder, FloatPlanWizard", font=small_font, fill=(217, 242, 245))
    draw.text((1472, 2440), "FloatPlanWizard.com", font=author_font, fill="white", anchor="ra")

    cover.convert("RGB").save(COVER_PATH, "JPEG", quality=94, optimize=True, progressive=True, dpi=(300, 300))
    return COVER_PATH


class CoverFlowable(Flowable):
    def __init__(self, path: Path, page_width: float, page_height: float, left: float, bottom: float, top: float):
        super().__init__()
        self.path = path
        self.width = page_width - left * 2
        self.height = page_height - bottom - top
        self.page_width = page_width
        self.page_height = page_height
        self.left = left
        self.bottom = bottom

    def draw(self):
        # Preserve the complete 5:8 cover inside the 2:3 PDF page; dark side bars blend into the artwork.
        draw_height = self.page_height
        draw_width = draw_height * 0.625
        self.canv.setFillColor(colors.HexColor("#031a2d"))
        self.canv.rect(-self.left, -self.bottom, self.page_width, self.page_height, stroke=0, fill=1)
        self.canv.drawImage(
            str(self.path), (self.page_width - draw_width) / 2 - self.left, -self.bottom,
            width=draw_width, height=draw_height, preserveAspectRatio=True, anchor="c", mask="auto",
        )


class PDFHeading(Paragraph):
    def __init__(self, text, style, anchor: str, level: int):
        super().__init__(f'<a name="{anchor}"/>{text}', style)
        self.fpw_anchor = anchor
        self.fpw_level = level
        self.fpw_title = re.sub(r"<[^>]+>", "", text)


class PDFChecklistItem(Flowable):
    """One wrapped checklist line with a printable, interactive AcroForm checkbox."""

    def __init__(self, markup: str, style: ParagraphStyle, field_name: str, tooltip: str):
        super().__init__()
        self.paragraph = Paragraph(markup, style)
        self.field_name = field_name
        self.tooltip = tooltip
        self.checkbox_size = 10
        self.gutter = 16
        self.spaceAfter = 2.5
        self._paragraph_width = 0
        self._paragraph_height = 0

    def wrap(self, avail_width, avail_height):
        self._paragraph_width = max(1, avail_width - self.gutter)
        _, self._paragraph_height = self.paragraph.wrap(self._paragraph_width, avail_height)
        self.width = avail_width
        self.height = max(self._paragraph_height, self.checkbox_size + 2)
        return self.width, self.height

    def draw(self):
        checkbox_y = max(0, self.height - self.checkbox_size - 1)
        self.canv.acroForm.checkbox(
            checked=False,
            buttonStyle="check",
            shape="square",
            fillColor=colors.white,
            borderColor=colors.HexColor("#087f8c"),
            textColor=colors.HexColor("#062743"),
            borderWidth=1,
            size=self.checkbox_size,
            x=0,
            y=checkbox_y,
            tooltip=self.tooltip,
            name=self.field_name,
            annotationFlags="print",
            fieldFlags="",
            forceBorder=True,
            relative=True,
        )
        self.paragraph.drawOn(self.canv, self.gutter, self.height - self._paragraph_height)


class BookDocTemplate(BaseDocTemplate):
    def __init__(self, filename, **kwargs):
        super().__init__(filename, **kwargs)
        frame = Frame(
            self.leftMargin, self.bottomMargin, self.width, self.height,
            id="book-frame", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0,
        )
        self.addPageTemplates(PageTemplate(id="book", frames=[frame], onPage=self.on_page))
        self._bookmark_count = 0

    def beforeDocument(self):
        self.canv.setTitle(FULL_TITLE)
        self.canv.setAuthor(AUTHOR)
        self.canv.setSubject(DESCRIPTION)
        self.canv.setKeywords(KEYWORDS)
        self.canv.setCreator("FloatPlanWizard reproducible publication builder")

    def on_page(self, canvas, doc):
        if doc.page <= 1:
            return
        canvas.saveState()
        canvas.setFont("FPWArial", 7.7)
        canvas.setFillColor(colors.HexColor("#526b7d"))
        canvas.drawString(doc.leftMargin, doc.pagesize[1] - 0.36 * inch, "FLOATPLANWIZARD  |  SOLO BOATING SAFETY")
        canvas.drawCentredString(doc.pagesize[0] / 2, 0.31 * inch, str(doc.page))
        canvas.restoreState()

    def afterFlowable(self, flowable):
        if isinstance(flowable, PDFHeading):
            key = flowable.fpw_anchor
            self.canv.bookmarkPage(key)
            self.canv.addOutlineEntry(flowable.fpw_title, key, level=flowable.fpw_level, closed=False)
            self.notify("TOCEntry", (flowable.fpw_level, flowable.fpw_title, self.page, key))


def register_pdf_fonts() -> None:
    pdfmetrics.registerFont(TTFont("FPWArial", str(ARIAL)))
    pdfmetrics.registerFont(TTFont("FPWArialBold", str(ARIAL_BOLD)))
    pdfmetrics.registerFont(TTFont("FPWArialItalic", str(ARIAL_ITALIC)))
    addMapping("FPWArial", 0, 0, "FPWArial")
    addMapping("FPWArial", 1, 0, "FPWArialBold")
    addMapping("FPWArial", 0, 1, "FPWArialItalic")


def pdf_styles():
    styles = getSampleStyleSheet()
    navy = colors.HexColor("#062743")
    ink = colors.HexColor("#173043")
    teal = colors.HexColor("#087f8c")
    base = dict(fontName="FPWArial", textColor=ink, allowWidows=0, allowOrphans=0)
    return {
        "body": ParagraphStyle("FPWBody", parent=styles["BodyText"], fontSize=9.4, leading=13.2,
                               spaceAfter=6.5, **base),
        "h1": ParagraphStyle("FPWH1", fontName="FPWArialBold", fontSize=20, leading=23,
                             textColor=navy, spaceBefore=9, spaceAfter=12, keepWithNext=True),
        "h2": ParagraphStyle("FPWH2", fontName="FPWArialBold", fontSize=13.2, leading=16,
                             textColor=teal, spaceBefore=11, spaceAfter=5.5, keepWithNext=True),
        "h3": ParagraphStyle("FPWH3", fontName="FPWArialBold", fontSize=10.7, leading=13,
                             textColor=navy, spaceBefore=8, spaceAfter=4, keepWithNext=True),
        "caption": ParagraphStyle("FPWCaption", fontName="FPWArial", fontSize=7.6, leading=10.2,
                                  textColor=colors.HexColor("#526b7d"), spaceAfter=8),
        "quote": ParagraphStyle("FPWQuote", fontName="FPWArialItalic", fontSize=9.2, leading=13,
                                textColor=ink, leftIndent=14, rightIndent=8, borderColor=colors.HexColor("#e6a93d"),
                                borderWidth=1, borderPadding=(12, 8, 12, 8), backColor=colors.HexColor("#f2f7f8"),
                                spaceBefore=18, spaceAfter=18),
        "list": ParagraphStyle("FPWList", fontName="FPWArial", fontSize=9.2, leading=12.7,
                               textColor=ink, leftIndent=4, spaceAfter=2.5),
        "checklist": ParagraphStyle("FPWChecklist", fontName="FPWArial", fontSize=9.2, leading=12.7,
                                     textColor=ink, leftIndent=0, firstLineIndent=0),
        "toc": ParagraphStyle("FPWTOC", fontName="FPWArial", fontSize=9.5, leading=13,
                              textColor=navy, leftIndent=12, firstLineIndent=-12, spaceBefore=3),
        "title": ParagraphStyle("FPWTitle", fontName="FPWArialBold", fontSize=25, leading=29,
                                textColor=navy, alignment=TA_CENTER, spaceAfter=10),
        "subtitle": ParagraphStyle("FPWSubtitle", fontName="FPWArial", fontSize=15, leading=19,
                                   textColor=teal, alignment=TA_CENTER, spaceAfter=18),
        "center": ParagraphStyle("FPWCenter", fontName="FPWArial", fontSize=10, leading=14,
                                 textColor=ink, alignment=TA_CENTER, spaceAfter=8),
    }


def add_pdf_element(element, story: list, styles: dict, depth: int = 0) -> None:
    tag = element.tag.lower() if isinstance(element.tag, str) else ""
    if tag == "section":
        is_chapter = "chapter" in (element.get("class") or "").split()
        if is_chapter and story:
            story.append(PageBreak())
        for child in element:
            add_pdf_element(child, story, styles, depth + 1)
        if element.get("id") == "title-page":
            story.append(PageBreak())
        return
    if tag in {"h1", "h2", "h3"}:
        level = {"h1": 0, "h2": 1, "h3": 2}[tag]
        anchor = element.getparent().get("id") if tag == "h1" else f"heading-{len(story)}"
        parent = element.getparent()
        if tag == "h1" and parent.get("id") == "title-page":
            story.append(Spacer(1, 0.65 * inch))
            story.append(Paragraph(inner_markup(element), styles["title"]))
        elif tag == "h1" and parent.get("data-nav") == "omit":
            story.append(Paragraph(inner_markup(element), styles[tag]))
        else:
            story.append(PDFHeading(inner_markup(element), styles[tag], anchor or f"heading-{len(story)}", level))
        return
    if tag == "p":
        classes = (element.get("class") or "").split()
        if "subtitle" in classes:
            style = styles["subtitle"]
        elif set(classes) & {"deck", "author", "publisher", "signature"}:
            style = styles["center"]
        else:
            style = styles["body"]
        story.append(Paragraph(inner_markup(element), style))
        return
    if tag in {"ul", "ol"}:
        is_checklist = "checklist" in (element.get("class") or "").split()
        if is_checklist:
            for item_index, li in enumerate(element.xpath("./li"), start=1):
                story.append(PDFChecklistItem(
                    inner_markup(li),
                    styles["checklist"],
                    checklist_field_name(element, item_index),
                    plain_text(li),
                ))
            story.append(Spacer(1, 5))
            return
        list_style = ParagraphStyle(
            f"FPWList-{len(story)}", parent=styles["list"], leftIndent=17, firstLineIndent=-14,
        )
        for index, li in enumerate(element.xpath("./li"), start=1):
            marker = f"{index}." if tag == "ol" else "&#8226;"
            story.append(Paragraph(f"{marker}&nbsp;&nbsp;{inner_markup(li)}", list_style))
        story.append(Spacer(1, 5))
        return
    if tag == "blockquote":
        story.append(Paragraph(inner_markup(element), styles["quote"]))
        return
    if tag == "figure":
        image_element = element.find("img")
        caption_element = element.find("figcaption")
        if image_element is None:
            return
        path = (MANUSCRIPT.parent / image_element.get("src")).resolve()
        with PILImage.open(path) as source:
            ratio = source.height / source.width
        width = 4.72 * inch
        graphic = Image(str(path), width=width, height=width * ratio)
        graphic.hAlign = "CENTER"
        caption = Paragraph(inner_markup(caption_element), styles["caption"]) if caption_element is not None else Spacer(1, 1)
        story.append(KeepTogether([graphic, Spacer(1, 3), caption]))
        return
    for child in element:
        add_pdf_element(child, story, styles, depth + 1)


def build_pdf(root) -> Path:
    register_pdf_fonts()
    styles = pdf_styles()
    page_width, page_height = 6 * inch, 9 * inch
    margin_x, bottom, top = 0.62 * inch, 0.58 * inch, 0.65 * inch
    doc = BookDocTemplate(
        str(PDF_PATH), pagesize=(page_width, page_height),
        leftMargin=margin_x, rightMargin=margin_x, topMargin=top, bottomMargin=bottom,
        title=FULL_TITLE, author=AUTHOR, subject=DESCRIPTION,
    )
    article = root.xpath('//article[@id="book"]')[0]
    story: list = [CoverFlowable(COVER_PATH, page_width, page_height, margin_x, bottom, top), PageBreak()]
    for opening in article.xpath("./section")[:2]:
        add_pdf_element(opening, story, styles)
    story.extend([PageBreak(), Paragraph("Contents", styles["h1"])])
    toc = TableOfContents()
    toc.levelStyles = [styles["toc"], ParagraphStyle("FPWTOC2", parent=styles["toc"], leftIndent=24, fontSize=8.5)]
    story.extend([toc, PageBreak()])
    for section in article.xpath("./section")[2:]:
        add_pdf_element(section, story, styles)
    doc.multiBuild(story)
    return PDF_PATH


def section_title(section) -> str:
    heading = section.find("h1")
    return " ".join(heading.itertext()).strip() if heading is not None else section.get("data-title", section.get("id", "Section"))


def xhtml_document(title: str, body_markup: str, *, body_type: str = "bodymatter") -> str:
    return f'''<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="en-US" xml:lang="en-US">
<head>
  <meta charset="utf-8"/>
  <title>{escape(title)}</title>
  <link rel="stylesheet" type="text/css" href="styles/epub.css"/>
</head>
<body epub:type="{body_type}">
{body_markup}
</body>
</html>'''


def epub_section_markup(section) -> str:
    clone = deepcopy(section)
    for image in clone.xpath(".//img"):
        image.set("src", f"images/{Path(image.get('src')).name}")
    for checklist in clone.xpath('.//ul[contains(concat(" ", normalize-space(@class), " "), " checklist ")]'):
        for item_index, item in enumerate(checklist.xpath("./li"), start=1):
            field_name = checklist_field_name(checklist, item_index)
            checkbox_id = f"checkbox-{field_name}"
            label_text = plain_text(item)
            original_text = item.text
            original_children = list(item)
            for child in original_children:
                item.remove(child)
            item.text = None

            checkbox = etree.Element("input")
            checkbox.set("type", "checkbox")
            checkbox.set("id", checkbox_id)
            checkbox.set("name", field_name)
            checkbox.set("class", "checklist-checkbox")
            checkbox.set("aria-label", label_text)

            print_box = etree.Element("span")
            print_box.set("class", "checklist-print-box")
            print_box.set("aria-hidden", "true")

            label = etree.Element("label")
            label.set("for", checkbox_id)
            label.text = original_text
            for child in original_children:
                label.append(child)

            item.extend([checkbox, print_box, label])
    return etree.tostring(clone, encoding="unicode", method="xml")


def build_epub(root) -> Path:
    article = root.xpath('//article[@id="book"]')[0]
    sections = article.xpath("./section")
    entries: list[dict[str, str]] = []
    for index, section in enumerate(sections, start=1):
        name = f"{index:02d}-{section.get('id')}.xhtml"
        entries.append({
            "id": f"section-{index}",
            "name": name,
            "title": section_title(section),
            "markup": epub_section_markup(section),
            "type": "frontmatter" if "front-matter" in (section.get("class") or "") else
                    "backmatter" if "back-matter" in (section.get("class") or "") else "bodymatter",
            "nav": section.get("data-nav", "include"),
        })

    nav_items = "\n".join(
        f'      <li><a href="{entry["name"]}">{escape(entry["title"])}</a></li>'
        for entry in entries if entry["nav"] != "omit"
    )
    nav_doc = f'''<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="en-US" xml:lang="en-US">
<head><meta charset="utf-8"/><title>Contents</title><link rel="stylesheet" type="text/css" href="styles/epub.css"/></head>
<body>
  <nav epub:type="toc" id="toc" aria-label="Table of contents">
    <h1>Contents</h1>
    <ol>
{nav_items}
    </ol>
  </nav>
  <nav epub:type="landmarks" aria-label="Guide">
    <h2>Guide</h2>
    <ol>
      <li><a epub:type="cover" href="cover.xhtml">Cover</a></li>
      <li><a epub:type="bodymatter" href="{next(e['name'] for e in entries if e['type'] == 'bodymatter')}">Start reading</a></li>
    </ol>
  </nav>
</body>
</html>'''

    cover_doc = '''<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="en-US" xml:lang="en-US">
<head><meta charset="utf-8"/><title>Cover</title><link rel="stylesheet" type="text/css" href="styles/epub.css"/></head>
<body epub:type="cover"><figure><img src="images/cover.jpg" alt="Cover of Solo Boating Safety: A Practical Guide from Kayaks to Cruisers."/></figure></body>
</html>'''

    image_paths = sorted(GRAPHICS.glob("*.png"))
    manifest_sections = "\n".join(
        f'    <item id="{entry["id"]}" href="{entry["name"]}" media-type="application/xhtml+xml"/>'
        for entry in entries
    )
    manifest_images = "\n".join(
        f'    <item id="image-{index}" href="images/{path.name}" media-type="image/png"/>'
        for index, path in enumerate(image_paths, start=1)
    )
    spine = "\n".join(f'    <itemref idref="{entry["id"]}"/>' for entry in entries)
    package = f'''<?xml version="1.0" encoding="utf-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="book-id" xml:lang="en-US" prefix="schema: http://schema.org/">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:identifier id="book-id">urn:uuid:1d8eac3e-894c-4f64-bab3-9158cdb20a2d</dc:identifier>
    <dc:title>{escape(FULL_TITLE)}</dc:title>
    <dc:creator id="creator">{escape(AUTHOR)}</dc:creator>
    <meta refines="#creator" property="role" scheme="marc:relators">aut</meta>
    <dc:publisher>{escape(PUBLISHER)}</dc:publisher>
    <dc:language>en-US</dc:language>
    <dc:description>{escape(DESCRIPTION)}</dc:description>
    <dc:rights>Copyright © 2026 FloatPlanWizard. All rights reserved.</dc:rights>
    <dc:date>2026</dc:date>
    <meta property="dcterms:modified">2026-08-13T00:00:00Z</meta>
    <meta property="schema:accessMode">textual</meta>
    <meta property="schema:accessMode">visual</meta>
    <meta property="schema:accessModeSufficient">textual</meta>
    <meta property="schema:accessibilityFeature">alternativeText</meta>
    <meta property="schema:accessibilityFeature">formControls</meta>
    <meta property="schema:accessibilityFeature">tableOfContents</meta>
    <meta property="schema:accessibilityFeature">structuralNavigation</meta>
    <meta property="schema:accessibilityHazard">none</meta>
    <meta property="schema:accessibilitySummary">Headings, structural navigation, a linked table of contents, meaningful link text, and text alternatives for informative figures are provided.</meta>
  </metadata>
  <manifest>
    <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
    <item id="cover-page" href="cover.xhtml" media-type="application/xhtml+xml"/>
    <item id="cover-image" href="images/cover.jpg" media-type="image/jpeg" properties="cover-image"/>
    <item id="css" href="styles/epub.css" media-type="text/css"/>
{manifest_sections}
{manifest_images}
  </manifest>
  <spine>
    <itemref idref="cover-page"/>
{spine}
  </spine>
</package>'''

    container = '''<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles><rootfile full-path="EPUB/package.opf" media-type="application/oebps-package+xml"/></rootfiles>
</container>'''

    with ZipFile(EPUB_PATH, "w") as archive:
        archive.writestr("mimetype", "application/epub+zip", compress_type=ZIP_STORED)
        archive.writestr("META-INF/container.xml", container, compress_type=ZIP_DEFLATED)
        archive.writestr("EPUB/package.opf", package, compress_type=ZIP_DEFLATED)
        archive.writestr("EPUB/nav.xhtml", nav_doc, compress_type=ZIP_DEFLATED)
        archive.writestr("EPUB/cover.xhtml", cover_doc, compress_type=ZIP_DEFLATED)
        archive.writestr("EPUB/styles/epub.css", EPUB_CSS.read_text(), compress_type=ZIP_DEFLATED)
        archive.write(COVER_PATH, "EPUB/images/cover.jpg", compress_type=ZIP_DEFLATED)
        for image_path in image_paths:
            archive.write(image_path, f"EPUB/images/{image_path.name}", compress_type=ZIP_DEFLATED)
        for entry in entries:
            archive.writestr(
                f'EPUB/{entry["name"]}',
                xhtml_document(entry["title"], entry["markup"], body_type=entry["type"]),
                compress_type=ZIP_DEFLATED,
            )
    return EPUB_PATH


def set_cell_or_paragraph_bottom_border(paragraph, color="087F8C", size="20") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_docx_hyperlink(paragraph, text: str, url: str, *, bold=False, italic=False):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "087F8C")
    props.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)
    if bold:
        props.append(OxmlElement("w:b"))
    if italic:
        props.append(OxmlElement("w:i"))
    run.append(props)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def append_docx_inline(paragraph, element, *, bold=False, italic=False) -> None:
    if element.text:
        run = paragraph.add_run(element.text)
        run.bold = bold
        run.italic = italic
    for child in element:
        tag = child.tag.lower() if isinstance(child.tag, str) else ""
        if tag == "a":
            add_docx_hyperlink(
                paragraph, "".join(child.itertext()), child.get("href", ""), bold=bold, italic=italic,
            )
        elif tag == "br":
            paragraph.add_run().add_break(WD_BREAK.LINE)
        else:
            append_docx_inline(
                paragraph,
                child,
                bold=bold or tag in {"strong", "b"},
                italic=italic or tag in {"em", "i"},
            )
        if child.tail:
            run = paragraph.add_run(child.tail)
            run.bold = bold
            run.italic = italic


def add_docx_checkbox(paragraph, field_name: str, label: str) -> None:
    """Add an unchecked Word checkbox content control with a visible print fallback glyph."""
    sdt = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")

    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), label)
    properties.append(alias)

    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), field_name)
    properties.append(tag)

    control_id = OxmlElement("w:id")
    control_id.set(qn("w:val"), str(crc32(field_name.encode("utf-8")) & 0x7FFFFFFF))
    properties.append(control_id)

    checkbox = OxmlElement("w14:checkbox")
    checked = OxmlElement("w14:checked")
    checked.set(qn("w14:val"), "0")
    checkbox.append(checked)
    checked_state = OxmlElement("w14:checkedState")
    checked_state.set(qn("w14:val"), "2612")
    checked_state.set(qn("w14:font"), "MS Gothic")
    checkbox.append(checked_state)
    unchecked_state = OxmlElement("w14:uncheckedState")
    unchecked_state.set(qn("w14:val"), "2610")
    unchecked_state.set(qn("w14:font"), "MS Gothic")
    checkbox.append(unchecked_state)
    properties.append(checkbox)
    sdt.append(properties)

    content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), "MS Gothic")
    run_properties.append(fonts)
    run.append(run_properties)
    text = OxmlElement("w:t")
    text.text = "☐"
    run.append(text)
    content.append(run)
    sdt.append(content)

    paragraph._p.append(sdt)
    paragraph.add_run(" ")


def set_picture_alt_text(inline_shape, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", "FloatPlanWizard safety figure")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_docx_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(23, 48, 67)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Title": (28, "062743", 0, 14),
        "Subtitle": (15, "087F8C", 0, 12),
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "062743", 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for style_name, (size, color, before, after) in specs.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Checklist" not in styles:
        checklist = styles.add_style("Checklist", WD_STYLE_TYPE.PARAGRAPH)
    else:
        checklist = styles["Checklist"]
    checklist.base_style = normal
    checklist.font.name = "Calibri"
    checklist.font.size = Pt(11)
    checklist.paragraph_format.left_indent = Inches(0.375)
    checklist.paragraph_format.first_line_indent = Inches(-0.188)
    checklist.paragraph_format.space_after = Pt(4)
    checklist.paragraph_format.line_spacing = 1.25

    if "FPW Quote" not in styles:
        quote = styles.add_style("FPW Quote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote = styles["FPW Quote"]
    quote.base_style = normal
    quote.font.italic = True
    quote.paragraph_format.left_indent = Inches(0.25)
    quote.paragraph_format.right_indent = Inches(0.15)
    quote.paragraph_format.space_before = Pt(18)
    quote.paragraph_format.space_after = Pt(18)

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor(82, 107, 125)
    caption.font.italic = False
    caption.paragraph_format.space_after = Pt(8)


def add_docx_element(document: Document, element) -> None:
    tag = element.tag.lower() if isinstance(element.tag, str) else ""
    if tag == "section":
        for child in element:
            add_docx_element(document, child)
        return
    if tag in {"h1", "h2", "h3"}:
        paragraph = document.add_paragraph(style={"h1": "Heading 1", "h2": "Heading 2", "h3": "Heading 3"}[tag])
        append_docx_inline(paragraph, element)
        return
    if tag == "p":
        paragraph = document.add_paragraph()
        classes = (element.get("class") or "").split()
        if set(classes) & {"author", "publisher", "signature", "subtitle"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        append_docx_inline(paragraph, element)
        return
    if tag in {"ul", "ol"}:
        checklist = "checklist" in (element.get("class") or "").split()
        for item_index, item in enumerate(element.xpath("./li"), start=1):
            paragraph = document.add_paragraph(style="Checklist" if checklist else ("List Number" if tag == "ol" else "List Bullet"))
            if checklist:
                add_docx_checkbox(
                    paragraph,
                    checklist_field_name(element, item_index),
                    plain_text(item),
                )
            append_docx_inline(paragraph, item)
        return
    if tag == "blockquote":
        paragraph = document.add_paragraph(style="FPW Quote")
        append_docx_inline(paragraph, element)
        return
    if tag == "figure":
        image_element = element.find("img")
        caption_element = element.find("figcaption")
        if image_element is None:
            return
        path = (MANUSCRIPT.parent / image_element.get("src")).resolve()
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = paragraph.add_run().add_picture(str(path), width=Inches(6.15))
        set_picture_alt_text(shape, image_element.get("alt", ""))
        if caption_element is not None:
            caption = document.add_paragraph(style="Caption")
            append_docx_inline(caption, caption_element)
        return
    for child in element:
        add_docx_element(document, child)


def add_editorial_cover(document: Document) -> None:
    label = document.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = label.add_run("FLOATPLANWIZARD SOLO BOATING SAFETY SERIES")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(8, 127, 140)
    label.paragraph_format.space_before = Pt(34)
    label.paragraph_format.space_after = Pt(68)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run(TITLE)
    title.paragraph_format.space_after = Pt(12)

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.add_run(SUBTITLE)
    subtitle.paragraph_format.space_after = Pt(22)
    set_cell_or_paragraph_bottom_border(subtitle)

    deck = document.add_paragraph()
    deck.paragraph_format.space_before = Pt(20)
    deck.paragraph_format.space_after = Pt(82)
    run = deck.add_run(DESCRIPTION)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(82, 107, 125)

    author = document.add_paragraph()
    author.add_run("Larry W.\n").bold = True
    author.add_run("Founder, FloatPlanWizard\n")
    link_holder = html.fromstring('<p><a href="https://floatplanwizard.com/">FloatPlanWizard.com</a></p>')
    append_docx_inline(author, link_holder)
    document.add_page_break()


def build_docx(root) -> Path:
    document = Document()
    properties = document.core_properties
    properties.title = FULL_TITLE
    properties.author = AUTHOR
    properties.subject = DESCRIPTION
    properties.keywords = KEYWORDS
    properties.comments = "Generated from the canonical FloatPlanWizard semantic HTML manuscript."
    properties.language = "en-US"

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_docx_styles(document)

    header = section.header.paragraphs[0]
    header.text = "FLOATPLANWIZARD  |  SOLO BOATING SAFETY"
    header.style = document.styles["Caption"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(section.footer.paragraphs[0])

    add_editorial_cover(document)
    contents = document.add_paragraph(style="Heading 1")
    contents.add_run("Contents")
    article = root.xpath('//article[@id="book"]')[0]
    for source_section in [s for s in article.xpath("./section") if s.get("data-nav") != "omit"]:
        entry = document.add_paragraph(style="List Bullet")
        entry.add_run(section_title(source_section))
    document.add_page_break()

    for source_section in article.xpath("./section"):
        if source_section.get("id") == "title-page":
            continue
        add_docx_element(document, source_section)

    document.save(DOCX_PATH)
    return DOCX_PATH


def copy_web_downloads() -> None:
    WEB_DOWNLOADS.mkdir(parents=True, exist_ok=True)
    shutil.copy2(PDF_PATH, WEB_DOWNLOADS / PDF_PATH.name)
    shutil.copy2(EPUB_PATH, WEB_DOWNLOADS / EPUB_PATH.name)


def build_all_publications() -> list[Path]:
    DIST.mkdir(parents=True, exist_ok=True)
    root = parse_manuscript()
    build_cover()
    outputs = [build_pdf(root), build_epub(root), build_docx(root), COVER_PATH]
    copy_web_downloads()
    return outputs


if __name__ == "__main__":
    outputs = build_all_publications()
    for output in outputs:
        print(output.relative_to(ROOT))
